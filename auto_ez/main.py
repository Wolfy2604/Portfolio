import docx, re, json, zipfile, shutil, os
from datetime import datetime
from docxtpl import DocxTemplate, RichText
from pprint import pprint
from validation import valid
from pytrovich.enums import NamePart, Gender, Case
from pytrovich.maker import PetrovichDeclinationMaker
from pytrovich.detector import PetrovichGenderDetector

print('HELLO')
def start(filename, input):
    print('START MAIN')
    maker = PetrovichDeclinationMaker()
    detector = PetrovichGenderDetector()
    table_lst = []
    alert = []
    dt = datetime.now()
    date1 = f'{dt.day if dt.day > 9 else f"0{dt.day}"}'
    date2 = f'{dt.month if dt.month > 9 else f"0{dt.month}"}'
    date3 = f'{dt.year if dt.year > 9 else f"0{dt.year}"}'
    date = f'{date1}.{date2}.{date3}'
    year = str(datetime.now().year)[2:4]
    main_data = {
        "num_ez": input["num_ez"],
        "today": input["ez_date"],
        "year": input["ez_date"][-2:],
        "bs": "",
        "bs2": "",
        "operator": "",
        "operator2": "",
        "oper_address": "",
        "address": "",
        "num_request": input["num_request"],
        "date_request": input["date_request"],
        "num_protocol": "",
        "date_protocol": "",
        "date_measure": 0,
        "source": [],
        "source2": [],
        "source_header": [],
        "source_merge": [[]],
        "source_merge2": [[]],
        "frequency": "",
        "frq_switcher": False,
        "result_header": [],
        "result0": [],
        "result1": [],
        "result2": [],
        "result3": [],
        "result4": [],
        "result5": [],
        "result_merge1": [],
        "result_merge2": [],
        "result_merge3": [],
        "result_merge4": [],
        "result_merge5": [],
        "devices": [],
        "devices_str": [],
        "temp": "",
        "humidity": 51,
        "pressure": 756,
        "repres": "",
        "repres_g": "",
        "specialist": "",
        "specialist_g": "",
        "expert": "",
        "customer": "",
        "mes_docs": [],
        "exp_docs": [],
        "is_plan": False
    }

    oper = {
        "МТС": ["ПАО «МТС»", "109147 г. Москва, ул. Марксистская, дом 4"],
        "МегаФон": ["ПАО «МегаФон»", "127006, г. Москва, Оружейный переулок, д. 41"],
        "ВымпелКом": ["ПАО «Вымпелком»", "127083,г. Москва, ул. 8 Марта, д. 10 стр. 14"],
        "К-телеком": ["ООО «К-телеком»", "350089, Краснодарский край, г. Краснодар, проспект Чекистов, 33/2"],
        "Миранда": ["ООО «Миранда-Медиа»", "295011, Республика Крым, г. Симферополь, ул. Героев Аджимушкая, д. 9"],
        "Т2 Мобайл": ["ООО «Т2 Мобайл»", "108811, г. Москва, Киевское шоссе 22-й "
                     "(п Московский) километр, домовладение 6, строение 1, этаж 5 комната 33"]
    }
    with open('bd.json', encoding='utf8') as j:
        main_data["expert"] = json.load(j)["expert"]
    print(main_data["expert"])

    '''DEL IMG'''
    shutil.rmtree('img/word/media')
    os.mkdir('img/word/media')

    try:
        '''TEXT'''
        start_doc = docx.Document(f'input/{filename}')
        parse_prot = ' '.join([p.text for p in start_doc.paragraphs])
        pprint(parse_prot)
        '''IMG'''
        doc = f'input/{filename}'
        with zipfile.ZipFile(doc) as zf:
            for name in zf.infolist():
                if name.filename.startswith('word/media/'):
                    zf.extract(name, 'img')

        '''PARSE'''
        num_prot_idx = parse_prot.find('№ ') + 2
        main_data["num_protocol"] = parse_prot[num_prot_idx:num_prot_idx + 4].split('/')[0]
        date_prot_idx = parse_prot.find('Дата оформления протокола:') + 26
        main_data["date_protocol"] = parse_prot[date_prot_idx:date_prot_idx + 11].strip('г ')
        date_measure_idx = parse_prot.find('Дата проведения измерений:') + 26
        main_data["date_measure"] = parse_prot[date_measure_idx:date_measure_idx + 11].strip('г ')

        temp_idx = parse_prot.find('воздуха, °C:') + 12
        temp_idx2 = parse_prot.find('Влажность')
        main_data["temp"] = parse_prot[temp_idx:temp_idx2 - 1].strip()
        hum_idx = parse_prot.find('воздуха, %:') + 11
        hum_idx2 = parse_prot.find('Атмосферное')
        main_data["humidity"] = parse_prot[hum_idx:hum_idx2 - 1].strip()
        prs_idx = parse_prot.find('мм.рт.ст.') + 9
        main_data["pressure"] = parse_prot[prs_idx:prs_idx + 5].strip(': Н')
        cus_idx = parse_prot.find('данные заказчика:') + 17
        cus_idx2 = parse_prot.find('ИНН') - 1
        main_data["customer"] = parse_prot[cus_idx:cus_idx2].strip(' ,')
        phrase = parse_prot[parse_prot.find('представителя организации:') + 26:parse_prot.find('Нормативные документы') - 1]
        if phrase.find('»') > 1:
            rep_idx = phrase.find('»') + 1
        else:
            rep_idx = 0
        if parse_prot.find('итуационный план') >= 0:
            main_data["is_plan"] = True

        main_data["repres"] = phrase[rep_idx:].strip(': ')
        mes_docs_idx = parse_prot.find('проводились измерения') + 22
        mes_docs_idx2 = parse_prot.find('основании которых проводилась оценка') - 26
        for idx, doc in enumerate(parse_prot[mes_docs_idx:mes_docs_idx2].split(';')):
            if idx + 1 == len(parse_prot[mes_docs_idx:mes_docs_idx2].split(';')):
                rt1 = RichText(f'{doc.strip(" .:")}', size=24)
                main_data["mes_docs"].append(rt1)
            else:
                rt2 = RichText(f'{doc.strip(" .:")}; ', size=24)
                main_data["mes_docs"].append(rt2)
        exp_docs_idx = parse_prot.find('проводилась оценка') + 19
        exp_docs_idx2 = parse_prot.find('Источники физических факторов') - 1
        for idx, doc in enumerate(parse_prot[exp_docs_idx:exp_docs_idx2].split(';')):
            if idx + 1 == len(parse_prot[exp_docs_idx:exp_docs_idx2].split(';')):
                rt3 = RichText(f'{doc.strip(" .:")}', size=24)
                main_data["exp_docs"].append(rt3)
            else:
                rt4 = RichText(f'{doc.strip(" .:")}; ', size=24)
                main_data["exp_docs"].append(rt4)

        '''TABLES'''
        for tbl in start_doc.tables:
            table = []
            for row in tbl.rows:
                t_row = [[cell.text for cell in row.cells]]
                table.append(t_row)
            table_lst.append(table)
        pprint(table_lst)

        def get_bs(string, string2=None):
            print('GET BS2: ', type(string2))
            zpt = string.find(',')
            word = string.find('принадлежащ')
            idx2 = zpt if zpt > 0 else word
            if string.find('№') and (zpt or word):
                print(idx2)
                main_data["bs"] = string[(string.find('№') + 1):idx2].strip(' ').replace('\n', '')
            else:
                alert.append('Нет номера БС!')
            if  main_data["bs"].find('БС') >=0:
                main_data["bs"].replace('БС', '')
            find = False
            for key, val in oper.items():
                if string.find(key) >= 0:
                    find = True
                    print('FIND', string.find(key))
                    main_data["operator"] = val[0]
                    main_data["oper_address"] = val[1].replace('\n', '')
                    break
            if not find:
                alert.append('Оператор № 1 не нашелся!')
            find2 = False
            if string2:
                zpt2 = string2.find(',')
                word2 = string2.find('принадлежащ')
                idx2_2 = zpt2 if zpt2 > 0 else word2
                if string2.find('№') and (zpt2 > 0 or word2 > 0):
                    main_data["bs2"] = string2[(string2.find('№') + 1):idx2_2].strip(' ').replace('\n', '')
                elif string2.find('№') and (zpt2 < 0 and word2 < 0):
                    if string2.find('ООО') > 0:
                        idx2_2 = string2.find('ООО') - 1
                    elif string2.find('ПАО') > 0:
                        idx2_2 = string2.find('ПАО') - 1
                    main_data["bs2"] = string2[(string2.find('№') + 1):idx2_2].strip()
                else:
                    alert.append('Нет номера БС 2!')
                for key, val in oper.items():
                    if string2.find(key) >= 0:
                        find2 = True
                        print('FIND', string2.find(key))
                        main_data["operator2"] = val[0]
                        break
                if not find2:
                    alert.append('Оператор № 2 не нашелся!')

        bs = table_lst[1][0][0][2]
        if bs.count('базовая станция') == 2:
            bs1 = bs.split('базовая станция')[1].strip(' .')
            bs2 = bs.split('базовая станция')[2].strip(' .')
            get_bs(bs1, bs2)
        elif bs.count('базовая станция') == 1:
            get_bs(bs)
        else:
            alert.append('Передатчик без номера')
            main_data["bs"] = bs.split(',')[0].strip()
            main_data["operator"] = bs.split(',')[1].strip('. ')

        if 'Фактический' in table_lst[1][1][0][0]:
            address = table_lst[1][1][0][1].strip().replace('\n', '')
        else:
            address = table_lst[1][1][0][0].strip().replace('\n', '')
        main_data["address"] = address

        tbl_shift = 0
        if len(table_lst[3]) > 1:
            main_data["frq_switcher"] = True
            tbl_shift = 1
            main_data["source_header"].append(table_lst[3][0])
            switch_src = False
            switch_src2 = False
            for src in table_lst[3][1:]:
                for idx, item in enumerate(src[0]):
                    # print(src[0][idx])
                    src[0][idx] = item.replace('\n', '').strip(' ,.')
                if src[0][0] == src[0][1]:
                    if not switch_src:
                        main_data["source_merge"][0].append([[src[0]]])
                        switch_src = True
                    else:
                        main_data["source_merge2"][0].append([[src[0]]])
                        switch_src2 = True
                else:
                    if not switch_src2:
                        main_data["source"].append(src)
                    else:
                        main_data["source2"].append(src)

        else:
            main_data["source"] = table_lst[3][0][0][1].replace('\n', '').strip(' ,.')
            frq = table_lst[4 - tbl_shift][0][0][1].replace('\n', '').strip(' ,.')
            main_data["frequency"] = frq

        res = table_lst[5 - tbl_shift][2:]

        '''RESULT'''
        counter = 1
        counter_mrg = 0
        main_data["result_header"] = [[[
            '№ п/п', 'Место проведения измерения/описание точек измерения',
            f'Плотность потока энергии (мкВт/см{chr(0x00B2)})', '(Up)*',
            f'ПДУ, (мкВт/см{chr(0x00B2)}) для населения'
        ]]]
        for idx, row in enumerate(res):
            if not row[0][2]:
                row[0][2] = 'менее 3'
            if not row[0][3]:
                row[0][3] = '-'
            if row[0][0] == row[0][1]:
                counter_mrg = counter_mrg + 1
                print(counter_mrg)
                main_data[f"result_merge{counter_mrg}"].append([row[0][0], '', '', '', ''])
                continue
            main_data[f"result{counter_mrg}"].append([counter, f'К.т.{counter} {row[0][1].replace('\n', '')}',
                                                      row[0][2], row[0][3], row[0][4]])
            counter += 1
            print('CHECKING')
            # if row[0][2].find('<') >= 0:
            #     continue
            # elif row[0][2] and idx > 1:
            #     if float(row[0][2].strip().replace(',', '.')) <= 10:
            #         continue
            #     else:
            #         alert.append('Превышение ППЭ!')


        print('RESULT READY')

        for idx, dvs in enumerate(table_lst[2]):
            if idx == 0:
                continue
            dv = [dvs[0][1].strip().replace('\n', ' '), dvs[0][2].strip().replace('\n', ' '),
                  dvs[0][4].strip().replace('\n', ' ')]
            d = dvs[0][3].strip('до ').strip('г.')
            dv.append(d)
            main_data["devices"].append(dv)
        main_data["specialist"] = table_lst[6 - tbl_shift][0][0][1].replace('\n', '').strip('_ ')
        if not input["expert"]:
            pass
        else:
            with open('bd.json', 'w',  encoding='utf-8') as bd:
                json.dump({"expert": input["expert"]}, bd)
            main_data["expert"] = input["expert"]
        pprint(main_data)

        '''GENITIVE'''
        def gen(string):
            try:
                lstname = re.search(r'[А-Яа-я]{3,15}', string.strip('. '))[0]
                dt = detector.detect(lastname=lstname)
                lstname_g = maker.make(NamePart.LASTNAME, dt, Case.GENITIVE, lstname)
                ini = re.search(r'[А-Я].[А-Я]', string.strip('. '))[0]
                return f'{lstname_g} {ini}'
            except:
                return f'{string}'

        def ins(string):
            try:
                lstname = re.search(r'[А-Яа-я]{3,15}', string.strip('. '))[0]
                dt = detector.detect(lastname=lstname)
                lstname_g = maker.make(NamePart.LASTNAME, dt, Case.INSTRUMENTAL, lstname)
                ini = re.search(r'[А-Я].[А-Я]', string.strip('. '))[0]
                return f'{lstname_g} {ini}'
            except:
                return f'{string}'

        main_data["repres_g"] = gen(main_data["repres"])
        main_data["specialist_g"] = ins(main_data["specialist"])

        '''CHECK'''
        dt_m = datetime.strptime(main_data["date_measure"], '%d.%m.%Y')
        for date in main_data["devices"]:
            dt_d = datetime.strptime(date[3].split('г')[0], '%d.%m.%Y')
            if dt_m > dt_d:
                alert.append('Сроки поверки истекли!')

    except ValueError:
        alert.append('Неверный формат файла!')
        return alert
    except Exception as e:
        print(repr(e))
        alert.append(repr(e))
        alert.append('Неверный формат файла!')
        return alert
    problems = valid(main_data)
    alert = alert + problems
    pprint(main_data)
    pprint(f"ALERT {alert}")

    '''DEVICES TO STR'''

    for dvs in main_data["devices"]:
        main_data["devices_str"].append(f'{dvs[0]}, заводской номер прибора № {dvs[1]},'
                                        f' свидетельство о поверке № {dvs[2]}, действительно до {dvs[3]} г., ')

    '''RENDER'''
    print('RENDER')
    print(main_data["result0"])
    final_doc = DocxTemplate("template/ЭЗ ИТОГОВЫЙ.docx")
    final_doc.render(main_data, autoescape=True)
    final_doc.save(f"output/ЭЗ{main_data["num_ez"]} - {main_data["bs"]}.docx")
    print('RENDER END')
    return alert


start('!!!26-940-_P3-42-53_.docx', {
    "num_request": 12,
    "date_request": '01.01.2003',
    "num_ez": 666,
    "expert": "",
    "ez_date": '02.01.2003'
    })


